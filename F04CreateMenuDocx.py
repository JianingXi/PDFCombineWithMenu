import os
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor
import natsort

def create_menu_and_index(image_dir, output_docx, initial_page_number=1):
    doc = Document()

    image_files = natsort.natsorted([f for f in os.listdir(image_dir) if f.endswith('.png')])
    for i, image_file in enumerate(image_files):
        doc.add_heading(f'Page {initial_page_number + i}', level=1)
        doc.add_paragraph(image_file)

    # 添加目录部分
    if image_dir:
        heading = doc.add_heading('附件目录', 0)

        # 设置标题的字体为黑体、居中、黑色、3号字
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.runs[0]
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(16)  # 3号字相当于16磅
        run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色

        # 设置文档的字体为宋体
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        page_number = initial_page_number
        for root, _, files in os.walk(image_dir):
            files = [f for f in files if f.lower().endswith('.jpg')]
            files = natsort.natsorted(files)
            for file in files:
                relative_path = os.path.relpath(root, image_dir)
                if relative_path == '.':
                    display_name = file  # 根目录的图片保持原名字不变
                else:
                    display_name = f"{relative_path.replace(os.sep, '_')}_{file}"
                    display_name = display_name.replace(' ', '_')  # 替换空格为下划线
                    display_name = display_name.replace('.', '_')  # 替换点为下划线，避免冲突
                    display_name += '.jpg'
                    display_name = display_name.replace('__jpg', '.jpg')  # 恢复正确的文件扩展名

                # 去掉文件名开头的4位数字
                display_name = '_'.join(display_name.split('_')[1:])

                # 自动省略文件名结尾的.jpg和_page_0001
                if display_name.lower().endswith('.jpg'):
                    display_name = display_name[:-4]

                # 使用正则表达式匹配并去掉_page_后跟四位数字的部分
                display_name = re.sub(r'_page_\d{4}$', '', display_name, flags=re.IGNORECASE)

                # 添加到docx文档
                p = doc.add_paragraph()
                run = p.add_run(display_name)

                # 设置字体为宋体和更小的大小
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(10)

                p.add_run().add_tab()
                tab_stop = p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

                p.add_run(str(page_number))

                # 设置段落对齐和间距
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = Pt(12)  # 行距设置为12磅

                page_number += 1

    # 保存docx文档
    doc.save(output_docx)
    print(f"Document saved at {output_docx}")

# 示例使用
image_dir = r'C:\Users\素材Temp\Page'
output_docx = r'C:\Users\素材Temp\目录.docx'
initial_page_number = 1  # 替换为你的初始页码

create_menu_and_index(image_dir, output_docx, initial_page_number)
